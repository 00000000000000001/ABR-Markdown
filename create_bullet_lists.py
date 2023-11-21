import copy
from docx.text.paragraph import Paragraph
import re
import utils
import test

# def insert_bullet_list(p):
#     p_new = copy.deepcopy(p)
#     p._p.addnext(p_new._p)

#     return False

# # def run(document):
# #     p = document.paragraphs[0]
# #     p_new = copy.deepcopy(p)
# #     p._p.addnext(p_new._p)
# #     # new_para = Paragraph(p_new, p._parent)
# #     # print(p.text)

# #     return True

# def is_item(text):
#     return re.match(r"\s*\*\*.*", text)

# def duplicate(p):
#     p_new = copy.deepcopy(p)
#     p._p.addnext(p_new._p)
#     return p_new

# def remove_runs_to_reverse(p, n):
#     i = len(p.runs) - 1
#     while i > n:
#         p._p.remove(p.runs[i]._r)
#         i -= 1

# def remove_runs_to(p, n):
#     i = 0
#     for i in range(n):
#         p._p.remove(p.runs[0]._r)

# def add_bullet_item(p, run):
#     p_new = utils.insert_paragraph_after(p, "", "List Bullet")
#     r_new = copy.deepcopy(run)._r
#     p_new._p.append(r_new)
#     return p_new

# def remove_run(run, p):
#     i = len(p.runs) - 1
#     while i >= 0:
#         if p.runs[i]._r == run._r:
#             print(run._r)
#             p._p.remove(p.runs[i]._r)
#             return run
#         i -= 1
#     return None

# def append_run(run, p):
#     r_copy = copy.deepcopy(run)._r
#     p._p.append(r_copy)

# def iterate_lines(run, p_last, p):
#     """
#     - gehe splitlines elementweise durch
#     - wann immer ein Element ein Item darstellt:
#         - entferne dieses Element aus splitlines
#         - mache daraus einen "List Bullet"-Paragraphen
#         - ersetzen den run-Text durch den Text der sich ergibt, wenn splitlines wieder zusammengesetzt wird,
#         wenn dieser nicht leer ist. Sonst: lösche den run.
#     """

#     lines = run.text.splitlines()
#     i = 0
#     while i < len(lines):
#         line = lines[i]
#         if is_item(line):
#             r_new = copy.deepcopy(run)._r
#             r_new.text = line[2:]
#             p_new = utils.insert_paragraph_after(p_last, "", "List Bullet")
#             p_new._p.append(r_new)
#             p_last = p_new
#             lines.remove(line)
#             run.text = utils.concate(lines)
#             bool = True
#             i -= 1
#         elif bool == True:
#             p_new = copy.deepcopy(p)
#             p_new.text = None
#             p_last._p.addnext(p_new._p)
#             p_last = p_new
#             run.text = run.text.strip()
#             remove_run(run, p)
#             append_run(run, p_last)
#             bool = False
#         i += 1
#     return p_last

# def iterate_runs(p):
#         i = 0
#         p_last = p
#         while i < len(p.runs):
#             run = p.runs[i]
#             p_last, bool = iterate_lines(run, p_last, p)
            # i += 1

# def run(document):
#     # j = 0
#     # while j < len(document.paragraphs):
#     #     iterate_runs(document.paragraphs[j])
#     #     j += 1

#     p = document.paragraphs[0]

    

#     exit()
#     # return True

def run(document):
    test.bullet(document)
    return True