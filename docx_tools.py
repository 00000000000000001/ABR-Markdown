from docx import Document
import copy

document = Document()

p_src = document.add_paragraph("")

p_src.add_run("abc").bold = True
p_src.add_run("defg").underline = True
p_src.add_run("hijkl").italic = True


def in_which_run_is(m, p):
    if m < 0 or m > len(p.text) - 1:
        return None
    i = 0
    sum = 0
    for run in p.runs:
        sum += len(run.text)
        if sum - 1 >= m:
            return i
        i += 1
    return None


assert in_which_run_is(0, p_src) == 0
assert in_which_run_is(1, p_src) == 0
assert in_which_run_is(2, p_src) == 0

assert in_which_run_is(3, p_src) == 1
assert in_which_run_is(4, p_src) == 1
assert in_which_run_is(5, p_src) == 1
assert in_which_run_is(6, p_src) == 1

assert in_which_run_is(7, p_src) == 2
assert in_which_run_is(8, p_src) == 2
assert in_which_run_is(9, p_src) == 2
assert in_which_run_is(10, p_src) == 2
assert in_which_run_is(11, p_src) == 2

assert in_which_run_is(12, p_src) == None
assert in_which_run_is(-1, p_src) == None


def at_which_position_in_its_run_is(m, p):
    if m < 0 or m > len(p.text) - 1:
        return None
    if m < len(p.runs[0].text):
        return m
    sum = 0
    i = 0
    while sum - 1 < m:
        sum += len(p.runs[i].text)
        i += 1
    k = (sum - 1) - m
    return len(p.runs[i - 1].text) - 1 - k


assert at_which_position_in_its_run_is(0, p_src) == 0
assert at_which_position_in_its_run_is(1, p_src) == 1
assert at_which_position_in_its_run_is(2, p_src) == 2

assert at_which_position_in_its_run_is(3, p_src) == 0
assert at_which_position_in_its_run_is(4, p_src) == 1
assert at_which_position_in_its_run_is(5, p_src) == 2
assert at_which_position_in_its_run_is(6, p_src) == 3

assert at_which_position_in_its_run_is(7, p_src) == 0
assert at_which_position_in_its_run_is(8, p_src) == 1
assert at_which_position_in_its_run_is(9, p_src) == 2
assert at_which_position_in_its_run_is(10, p_src) == 3
assert at_which_position_in_its_run_is(11, p_src) == 4

assert at_which_position_in_its_run_is(12, p_src) == None
assert at_which_position_in_its_run_is(-1, p_src) == None

p_dest = document.add_paragraph("")


def cp(m, n, p_src, p_dest):
    r_start = in_which_run_is(m, p_src)
    r_finish = in_which_run_is(n, p_src)

    for i in range(r_start, r_finish + 1):
        run = p_src.runs[i]
        r_copy = copy.deepcopy(run)._r

        a = 0
        o = len(run.text)
        if i == r_start:
            a = at_which_position_in_its_run_is(m, p_src)
        if i == r_finish:
            o = at_which_position_in_its_run_is(n, p_src) + 1

        r_copy.text = r_copy.text[a:o]
        p_dest._p.append(r_copy)


cp(5, 7, p_src, p_dest)
assert p_dest.text == "fgh"
p_dest._p.clear()

cp(0, 0, p_src, p_dest)
assert p_dest.text == "a"
p_dest._p.clear()

cp(0, 2, p_src, p_dest)
assert p_dest.text == "abc"
p_dest._p.clear()

cp(0, 3, p_src, p_dest)
assert p_dest.text == "abcd"
p_dest._p.clear()

cp(0, 11, p_src, p_dest)
assert p_dest.text == "abcdefghijkl"
p_dest._p.clear()

cp(7, 11, p_src, p_dest)
assert p_dest.text == "hijkl"
p_dest._p.clear()

cp(11, 11, p_src, p_dest)
assert p_dest.text == "l"
p_dest._p.clear()

cp(0, 11, p_src, p_dest)


def remove_run(run, p):
    i = len(p.runs) - 1
    while i >= 0:
        if p.runs[i]._r == run._r:
            p._p.remove(p.runs[i]._r)
            return run
        i -= 1
    return None


def rm(m, n, p):
    if m < 0 or n > len(p.text) - 1:
        return

    r_start = in_which_run_is(m, p)
    r_finish = in_which_run_is(n, p)

    a = -1
    o = -1
    arr = []

    for i in range(r_start, r_finish + 1):
        run = p.runs[i]

        if i == r_start:
            a = at_which_position_in_its_run_is(m, p)
        if i == r_finish:
            o = at_which_position_in_its_run_is(n, p)
        if i > r_start and i < r_finish:
            arr.append(run)

    if r_start == r_finish:
        p.runs[r_start].text = p.runs[r_start].text[:a] + p.runs[r_finish].text[o + 1 :]
    else:
        p.runs[r_start].text = p.runs[r_start].text[:a]
        p.runs[r_finish].text = p.runs[r_finish].text[o + 1 :]
    for run in reversed(arr):
        remove_run(run, p)


rm(0, 0, p_dest)
assert p_dest.text == "bcdefghijkl"

p_dest._p.clear()
cp(0, 11, p_src, p_dest)
rm(5, 7, p_dest)
assert p_dest.text == "abcdeijkl"

p_dest._p.clear()
cp(0, 11, p_src, p_dest)
rm(0, 2, p_dest)
assert p_dest.text == "defghijkl"

p_dest._p.clear()
cp(0, 11, p_src, p_dest)
rm(0, 11, p_dest)
assert p_dest.text == ""

p_dest._p.clear()
cp(0, 11, p_src, p_dest)
rm(1, 1, p_dest)
assert p_dest.text == "acdefghijkl"

p_dest._p.clear()
cp(0, 11, p_src, p_dest)
rm(1, 12, p_dest)
assert p_dest.text == "abcdefghijkl"

p_dest._p.clear()
cp(0, 11, p_src, p_dest)
rm(0, 2, p_dest)
assert p_dest.text == "defghijkl"

p_dest._p.clear()
cp(0, 11, p_src, p_dest)
rm(3, 6, p_dest)
assert p_dest.text == "abchijkl"


def mv(m, n, p_src, p_dest):
    cp(m, n, p_src, p_dest)
    rm(m, n, p_src)


p_dest._p.clear()
mv(0, 0, p_src, p_dest)
assert p_src.text == "bcdefghijkl"
assert p_dest.text == "a"

p_src._p.clear()
p_src.add_run("abc").bold = True
p_src.add_run("defg").underline = True
p_src.add_run("hijkl").italic = True
p_dest._p.clear()
mv(0, 11, p_src, p_dest)
assert p_src.text == ""
assert p_dest.text == "abcdefghijkl"

p_src._p.clear()
p_src.add_run("abc").bold = True
p_src.add_run("defg").underline = True
p_src.add_run("hijkl").italic = True
p_dest._p.clear()
mv(3, 6, p_src, p_dest)
assert p_src.text == "abchijkl"
assert p_dest.text == "defg"


document.save("test.docx")
