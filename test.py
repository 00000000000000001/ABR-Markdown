from docx import Document
from docx_tools import cp, rm, mv
import utils
import copy
from docx.text.paragraph import Paragraph
from docx.oxml.shared import OxmlElement


# document = Document()

# p = document.add_paragraph("")

# p.add_run("Befunde:\n").underline = True
# p.add_run("**BEF1\n**BEF2\n")
# p.add_run("Procedere:\n").underline = True
# p.add_run("**PROC1\n**PROC2")

def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def do_stuff(n, p):
    line = ""
    i = n - 1
    while i < len(p.text) and p.text[i] != "\n":
        line += p.text[i]
        i += 1
    p_bullet = utils.insert_paragraph_after(p_last, "", "List Bullet")
    mv(n - 1, i - 1, p, p_bullet)
    if n == 0:
        rm(n - 1, n - 2, p)
    else:
        rm(n - 2, n - 2, p)
    print(line) 
    n -= 1
    return [n, p_bullet]

def bullet(document):
    l = 0
    while l < len(document.paragraphs):
        p = document.paragraphs[l]
        p_last = p
        is_new_line = True
        could_be_a_list_item = False
        moved = False
        i = 0
        while i < len(p.text):

            if p.text[i] == "\n":
                is_new_line = True
                i += 1
                continue

            if is_new_line and p.text[i] == "*":
                could_be_a_list_item = True
                is_new_line = False
                moved = False
                i += 1
                continue

            if could_be_a_list_item and p.text[i] == "*":

                if i > 1:
                    rm(i-2, i, p)
                    i -= 2
                else:
                    rm(i-1, i, p) 
                    i -= 1


                line = ""
                k = i
                while k < len(p.text) and p.text[k] != "\n":
                    line += p.text[k]
                    k += 1

                print(line)

                p_new = insert_paragraph_after(p_last, "", "List Bullet")
                mv(i,k - 1,p,p_new)
                p_last = p_new
                moved = True
                continue

            if moved and p.text[i] != "*":
                p_new = copy.deepcopy(p)
                p_new._p.clear()
                p_last._p.addnext(p_new._p)
                mv(i, len(p.text) - 1, p, p_new)
                if p.text[len(p.text) - 1] == "\n":
                    rm(len(p.text) - 1, len(p.text) - 1, p)
                p_last = p_new
                moved = False
                i += 1
                continue

            i += 1
        if p.text == "":
            utils.delete_paragraph(p)
        l += 1


# bullet(document)

# document.save("test.docx")
