import sys

sys.path.append("../src")
sys.path.append("./src")

from docx import Document
from create_bullet_lists import bulletList, blI


def run():

    # I: (** v)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** ")
    bulletList(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == ""

    # I: (** v)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** v")
    bulletList(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "v"

    # II: (** v\nw)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** ")
    p.add_run("\n")
    bulletList(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == ""
    assert document.paragraphs[1].text == ""

    # II: (** v\nw)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** v")
    p.add_run("\nw")
    bulletList(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "v"
    assert document.paragraphs[1].text == "w"

    # III: (\n** )
    document = Document()
    p = document.add_paragraph("")
    p.add_run("\n")
    p.add_run("** ")
    bulletList(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == ""
    assert document.paragraphs[1].text == ""

    # III: (w\n** v)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("w\n")
    p.add_run("** v")
    bulletList(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "w"
    assert document.paragraphs[1].text == "v"

    # IV: (w\n** v1\nv2)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("\n")
    p.add_run("** \n")
    bulletList(document)
    assert len(document.paragraphs) == 3
    assert document.paragraphs[0].text == ""
    assert document.paragraphs[1].text == ""
    assert document.paragraphs[2].text == ""

    # IV: (w\n** v1\nv2)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("w\n")
    p.add_run("** v1\nv2")
    bulletList(document)
    assert len(document.paragraphs) == 3
    assert document.paragraphs[0].text == "w"
    assert document.paragraphs[1].text == "v1"
    assert document.paragraphs[2].text == "v2"

    ##

    document = Document()
    document.add_paragraph("")
    bulletList(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == ""

    document = Document()
    document.add_paragraph("")
    document.add_paragraph("")
    bulletList(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == ""
    assert document.paragraphs[1].text == ""

    document = Document()
    p = document.add_paragraph("")
    p.add_run("")
    bulletList(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == ""

    document = Document()
    p = document.add_paragraph("")
    p.add_run("")
    p.add_run("")
    bulletList(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == ""

    # FOOBAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO")
    p.add_run("BAR")
    bulletList(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOOBAR"

    # FOO
    # BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n")
    p.add_run("BAR")
    bulletList(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO\nBAR"

    # **FOOBAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO")
    p.add_run("BAR")
    bulletList(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOOBAR"

    # **FOO**BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO")
    p.add_run("** BAR")
    bulletList(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO** BAR"

    # **FOO
    # **BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO\n")
    p.add_run("** BAR")
    bulletList(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    # **FOO
    # **BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO")
    p.add_run("\n** BAR")
    bulletList(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    # **FOO
    # **BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO\n** BAR")
    bulletList(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    # **FOO
    # BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO\nBAR")
    bulletList(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    # FOO
    # **BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n** BAR")
    bulletList(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    # FOO
    # **BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n** BAR")
    bulletList(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    # FOO
    # **BAR
    # ZIG
    # **ZAG
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n** BAR\nZIG\n** ZAG")
    bulletList(document)
    assert len(document.paragraphs) == 4
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"
    assert document.paragraphs[2].text == "ZIG"
    assert document.paragraphs[3].text == "ZAG"

    # **BAR
    # ZIG
    # **ZAG
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** BAR\nZIG\n** ZAG")
    bulletList(document)
    assert len(document.paragraphs) == 3
    assert document.paragraphs[0].text == "BAR"
    assert document.paragraphs[1].text == "ZIG"
    assert document.paragraphs[2].text == "ZAG"

    # Does it remove leading white
    # ** Foo
    # **  BAR
    # **
    # **
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO\n** BAR\n** \n** \n** asd ")
    bulletList(document)
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"
    assert document.paragraphs[2].text == ""
    assert document.paragraphs[3].text == ""
    assert document.paragraphs[4].text == "asd "

    # document.save("./test.docx")

    return
