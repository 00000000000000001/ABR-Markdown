import sys

sys.path.append("../src")
from docx import Document
import remove_comments, split_paragraph, create_bullet_lists

def run():
    doc = Document()
    p_src = doc.add_paragraph("")
    assert remove_comments.comments(doc) | split_paragraph.para(doc) | create_bullet_lists.bulletList(doc) == False

    doc = Document()
    p_src = doc.add_paragraph("{}")
    assert remove_comments.comments(doc) | split_paragraph.para(doc) | create_bullet_lists.bulletList(doc) == True

    doc = Document()
    p_src = doc.add_paragraph("FOO\n\nBAR")
    assert remove_comments.comments(doc) | split_paragraph.para(doc) | create_bullet_lists.bulletList(doc) == True

    doc = Document()
    p_src = doc.add_paragraph("** FOO")
    assert remove_comments.comments(doc) | split_paragraph.para(doc) | create_bullet_lists.bulletList(doc) == True

    doc = Document()
    p_src = doc.add_paragraph("FOO")
    assert remove_comments.comments(doc) | split_paragraph.para(doc) | create_bullet_lists.bulletList(doc) == False

    # doc.save("test.docx")
