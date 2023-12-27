import sys

sys.path.append("./src")

from docx import Document
from preconvert import preconvert

document = Document()
p = document.add_paragraph("")
p.add_run("** FOO")
preconvert(document)
assert document.paragraphs[0].text == "**FOO"