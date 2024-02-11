import sys

sys.path.append("../src")

from docx import Document
from remove_comments import comments

def run():
    document = Document()
    p = document.add_paragraph("")
    p.add_run("")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == ""

    # FOO
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO"

    # {Test}
    document = Document()
    p = document.add_paragraph("")
    p.add_run("{Test}")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == ""

    # {Test}FOO
    document = Document()
    p = document.add_paragraph("")
    p.add_run("{Test}FOO")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO"

    # FOO{Test}
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO{Test}")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO"

    # {Test}FOO{Test}
    document = Document()
    p = document.add_paragraph("")
    p.add_run("{Test}FOO{Test}")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO"

    # FOO
    # {Test}
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n{Test}")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO\n"

    # {Test}
    # FOO
    document = Document()
    p = document.add_paragraph("")
    p.add_run("{Test}\nFOO")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO"

    # {Test}
    #
    # FOO
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n\n{Test}")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO\n\n"

    #
    # {Test}
    # FOO
    document = Document()
    p = document.add_paragraph("")
    p.add_run("\nFOO\n{Test}")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "\nFOO\n"

    # {Test}
    # {Test}
    # FOO
    document = Document()
    p = document.add_paragraph("")
    p.add_run("{Test}\n{Test}\nFOO")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO"

    # {Test}{Test}
    # FOO
    document = Document()
    p = document.add_paragraph("")
    p.add_run("{Test}{Test}\nFOO")
    comments(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO"


    # document.save("./test.docx")
