import sys

sys.path.append("../src")

from docx import Document
from split_paragraph import para

def run():
    document = Document()
    p = document.add_paragraph("")
    p.add_run("")
    para(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == ""

    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n\nBAR")
    para(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n\n\nBAR")
    para(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "\nBAR"

    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n\n\n\nBAR")
    para(document)
    assert len(document.paragraphs) == 3
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == ""
    assert document.paragraphs[2].text == "BAR"

    # document.save("./test.docx")
