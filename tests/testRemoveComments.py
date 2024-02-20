import sys

sys.path.append("../src")
sys.path.append("./src")

from docx import Document
from comments import removeComments


def run():

    doc = Document()
    p = doc.add_paragraph("")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == ""

    doc = Document()
    p = doc.add_paragraph("\n")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "\n"

    doc = Document()
    p = doc.add_paragraph("{}")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == ""

    doc = Document()
    p = doc.add_paragraph("{}{}")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == ""

    doc = Document()
    p = doc.add_paragraph("{{}}")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "}"

    doc = Document()
    p = doc.add_paragraph("foo")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "foo"

    doc = Document()
    p = doc.add_paragraph("{foo}")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == ""

    doc = Document()
    p = doc.add_paragraph("{foo}{foo}")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == ""

    # doc = Document()
    # p = doc.add_paragraph("{fo{foo}o}")
    # removeComments(doc)
    # assert len(doc.paragraphs) == 1
    # assert doc.paragraphs[0].text == ""

    doc = Document()
    p = doc.add_paragraph("{foo}bar")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "bar"

    doc = Document()
    p = doc.add_paragraph("foo{bar}")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "foo"

    doc = Document()
    p = doc.add_paragraph("foo{bar}foo")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "foofoo"

    doc = Document()
    p = doc.add_paragraph("{foo}\nbar")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "bar"

    doc = Document()
    p = doc.add_paragraph("{foo}\n\nbar")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "bar"

    doc = Document()
    p = doc.add_paragraph("foo{bar}\n{foo}bar")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "foo\nbar"

    doc = Document()
    p = doc.add_paragraph("foo{bar}\n\n{foo}bar")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "foo\n\nbar"

    doc = Document()
    p = doc.add_paragraph("foo\n\n{bar}\n\nfoo")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "foo\n\nfoo"

    doc = Document()
    p = doc.add_paragraph("foo{bar}\n\n{bar}\n\nfoo")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "foo\n\nfoo"

    doc = Document()
    p = doc.add_paragraph("foo{bar}\n\n{bar}{bar}\n\nfoo")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "foo\n\nfoo"

    doc = Document()
    p = doc.add_paragraph("foo{bar}\n\n{bar}\n\nfoo\n{bar}")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "foo\n\nfoo\n"

    doc = Document()
    p = doc.add_paragraph("FOO{BAR")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "FOO{BAR"

    doc = Document()
    p = doc.add_paragraph("FOO}BAR")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "FOO}BAR"

    doc = Document()
    p = doc.add_paragraph("FOO}{BAR")
    removeComments(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "FOO}{BAR"
