import sys

sys.path.append("../src")
sys.path.append("./src")

from docx import Document
from bulletList import substitute


def test_replaceDoubleAsterisks():

    # I: (** v)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** ")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == ""

    # I: (** v)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** v")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "v"

    # II: (** v\nw)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** ")
    p.add_run("\n")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "\n"

    # II: (** v\nw)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** v")
    p.add_run("\nw")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "v\nw"

    # III: (\n** )
    document = Document()
    p = document.add_paragraph("")
    p.add_run("\n")
    p.add_run("** ")
    substitute(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == ""
    assert document.paragraphs[1].text == ""

    # III: (w\n** v)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("w\n")
    p.add_run("** v")
    substitute(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "w"
    assert document.paragraphs[1].text == "v"

    # IV: (w\n** v1\nv2)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("\n")
    p.add_run("** \n")
    substitute(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == ""
    assert document.paragraphs[1].text == "\n"


    # IV: (w\n** v1\nv2)
    document = Document()
    p = document.add_paragraph("")
    p.add_run("w\n")
    p.add_run("** v1\nv2")
    substitute(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "w"
    assert document.paragraphs[1].text == "v1\nv2"

    ##

    document = Document()
    document.add_paragraph("")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == ""

    document = Document()
    document.add_paragraph("")
    document.add_paragraph("")
    substitute(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == ""
    assert document.paragraphs[1].text == ""

    document = Document()
    p = document.add_paragraph("")
    p.add_run("")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == ""

    document = Document()
    p = document.add_paragraph("")
    p.add_run("")
    p.add_run("")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == ""

    # FOOBAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO")
    p.add_run("BAR")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOOBAR"

    # FOO
    # BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n")
    p.add_run("BAR")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO\nBAR"

    # **FOOBAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO")
    p.add_run("BAR")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOOBAR"

    # **FOO**BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO")
    p.add_run("** BAR")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO** BAR"

    # **FOO
    # **BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO\n")
    p.add_run("** BAR")
    substitute(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    # **FOO
    # **BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO")
    p.add_run("\n** BAR")
    substitute(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    # **FOO
    # **BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO\n** BAR")
    substitute(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    # **FOO
    # BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO\nBAR")
    substitute(document)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "FOO\nBAR"

    # FOO
    # **BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n** BAR")
    substitute(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    # FOO
    # **BAR
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n** BAR")
    substitute(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"

    # FOO
    # **BAR
    # ZIG
    # **ZAG
    document = Document()
    p = document.add_paragraph("")
    p.add_run("FOO\n** BAR\nZIG\n** ZAG")
    substitute(document)
    assert len(document.paragraphs) == 3
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR\nZIG"
    assert document.paragraphs[2].text == "ZAG"

    # **BAR
    # ZIG
    # **ZAG
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** BAR\nZIG\n** ZAG")
    substitute(document)
    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "BAR\nZIG"
    assert document.paragraphs[1].text == "ZAG"

    # Does it remove leading white
    # ** Foo
    # **  BAR
    # **
    # **
    document = Document()
    p = document.add_paragraph("")
    p.add_run("** FOO\n** BAR\n** \n** \n** asd ")
    substitute(document)
    assert document.paragraphs[0].text == "FOO"
    assert document.paragraphs[1].text == "BAR"
    assert document.paragraphs[2].text == ""
    assert document.paragraphs[3].text == ""
    assert document.paragraphs[4].text == "asd "

    return
