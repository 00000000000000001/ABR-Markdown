import sys

sys.path.append("../src")
sys.path.append("./src")
from docx import Document
import comments, paragraph, bulletList
from converterFunktions import checkAndConvert

def test_wasEdited():
    doc = Document()
    p_src = doc.add_paragraph("")
    assert comments.removeComments(doc) | paragraph.subdivide(doc) | bulletList.substitute(doc) == False

    doc = Document()
    p_src = doc.add_paragraph("{}")
    assert comments.removeComments(doc) | paragraph.subdivide(doc) | bulletList.substitute(doc) == True

    doc = Document()
    p_src = doc.add_paragraph("FOO\n\nBAR")
    assert comments.removeComments(doc) | paragraph.subdivide(doc) | bulletList.substitute(doc) == True

    doc = Document()
    p_src = doc.add_paragraph("** FOO")
    assert comments.removeComments(doc) | paragraph.subdivide(doc) | bulletList.substitute(doc) == True

    doc = Document()
    p_src = doc.add_paragraph("FOO")
    assert comments.removeComments(doc) | paragraph.subdivide(doc) | bulletList.substitute(doc) == False

    # doc.save("test.docx")

def test_templates_with_commands_are_not_converted():
    doc = Document()
    doc.add_paragraph("P1\n\nP2")
    doc.add_paragraph("** Element")
    doc.add_paragraph("{Kommentar}")
    doc.add_paragraph("$[datum]$")
    assert checkAndConvert(doc) == None

def test_templates_without_commands_are_converted():
    doc = Document()
    doc.add_paragraph("P1\n\nP2")
    doc.add_paragraph("** Element")
    doc.add_paragraph("{Kommentar}")
    assert checkAndConvert(doc) != None


