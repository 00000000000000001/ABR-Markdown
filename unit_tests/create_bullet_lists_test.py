import sys

sys.path.append("../ABR-Markdown")

from docx import Document
from create_bullet_lists import run

document = Document()
document.add_paragraph("")
run(document)
assert len(document.paragraphs) == 1

document = Document()
document.add_paragraph("")
document.add_paragraph("")
run(document)
assert len(document.paragraphs) == 2

document = Document()
p = document.add_paragraph("")
p.add_run("")
run(document)
assert len(document.paragraphs) == 1

document = Document()
p = document.add_paragraph("")
p.add_run("")
p.add_run("")
run(document)
assert len(document.paragraphs) == 1

# FOOBAR
document = Document()
p = document.add_paragraph("")
p.add_run("FOO")
p.add_run("BAR")
run(document)
assert len(document.paragraphs) == 1

# FOO
# BAR
document = Document()
p = document.add_paragraph("")
p.add_run("FOO\n")
p.add_run("BAR")
run(document)
assert len(document.paragraphs) == 1

# **FOOBAR
document = Document()
p = document.add_paragraph("")
p.add_run("**FOO")
p.add_run("BAR")
run(document)
assert len(document.paragraphs) == 1

# **FOO**BAR
document = Document()
p = document.add_paragraph("")
p.add_run("**FOO")
p.add_run("**BAR")
run(document)
assert len(document.paragraphs) == 1

# **FOO
# **BAR
document = Document()
p = document.add_paragraph("")
p.add_run("**FOO\n")
p.add_run("**BAR")
run(document)
assert len(document.paragraphs) == 2

# **FOO
# **BAR
document = Document()
p = document.add_paragraph("")
p.add_run("**FOO")
p.add_run("\n**BAR")
run(document)
assert len(document.paragraphs) == 2

# **FOO
# **BAR
document = Document()
p = document.add_paragraph("")
p.add_run("**FOO\n**BAR")
run(document)
assert len(document.paragraphs) == 2

# **FOO
# BAR
document = Document()
p = document.add_paragraph("")
p.add_run("**FOO\nBAR")
run(document)
assert len(document.paragraphs) == 2

# FOO
# **BAR
document = Document()
p = document.add_paragraph("")
p.add_run("FOO\n**BAR")
run(document)
assert len(document.paragraphs) == 2

# FOO
# **BAR
document = Document()
p = document.add_paragraph("")
p.add_run("FOO\n**BAR")
run(document)
assert len(document.paragraphs) == 2

# FOO
# **BAR
# ZIG
# **ZAG
document = Document()
p = document.add_paragraph("")
p.add_run("FOO\n**BAR\nZig\n**ZAG")
run(document)
assert len(document.paragraphs) == 4

# **BAR
# ZIG
# **ZAG
document = Document()
p = document.add_paragraph("")
p.add_run("**BAR\nZig\n**ZAG")
run(document)
assert len(document.paragraphs) == 3


document.save("./test.docx")
