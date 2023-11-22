import sys

sys.path.append("../ABR-Markdown")
from docx import Document
import remove_comments, split_paragraph, create_bullet_lists

doc = Document()
p_src = doc.add_paragraph("")
assert remove_comments.run(doc) | split_paragraph.run(doc) | create_bullet_lists.run(doc) == False

doc = Document()
p_src = doc.add_paragraph("{}")
assert remove_comments.run(doc) | split_paragraph.run(doc) | create_bullet_lists.run(doc) == True

doc = Document()
p_src = doc.add_paragraph("FOO\n\nBAR")
assert remove_comments.run(doc) | split_paragraph.run(doc) | create_bullet_lists.run(doc) == True

doc = Document()
p_src = doc.add_paragraph("**FOO")
assert remove_comments.run(doc) | split_paragraph.run(doc) | create_bullet_lists.run(doc) == True

doc = Document()
p_src = doc.add_paragraph("FOO")
assert remove_comments.run(doc) | split_paragraph.run(doc) | create_bullet_lists.run(doc) == False

doc.save("test.docx")
