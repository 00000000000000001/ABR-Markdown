import hashlib
import config
import sqlite3
from docx.text.paragraph import Paragraph
from docx.oxml.shared import OxmlElement
import copy


def get_hash(file):
    with file:
        bytes = file.read()
        hash = hashlib.md5(bytes).hexdigest()
    return hash


def is_registered(hash):
    con = sqlite3.connect(config.SQLITE_FILE_NAME)
    cur = con.cursor()
    cur.execute("CREATE TABLE IF NOT EXISTS hashes(hash)")
    res = cur.execute(f"SELECT hash FROM hashes WHERE hash = '{hash}'")
    is_registered = res.fetchone() is not None
    con.close()
    return is_registered


def getText(doc):
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return "\n".join(fullText)


def register_hash(hash):
    conn = sqlite3.connect(config.SQLITE_FILE_NAME)
    cur = conn.cursor()
    cur.execute(f"INSERT INTO hashes (hash) VALUES ('{hash}')")
    conn.commit()
    conn.close()

def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

def concate(arr):
    res = ''
    for line in arr:
        res += line + '\n'
    return res[:-1]

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

# def copy_to_new_paragraph(p, text, run, style=None):
#     new_p = insert_paragraph_after(p, "", style)
#     runner = new_p.add_run(text)
#     runner.bold = run.bold
#     runner.italic = run.italic
#     runner.underline = run.underline
#     runner.font.color.rgb = run.font.color.rgb
#     runner.font.name = run.font.name
#     runner.font.size = run.font.size
#     return new_p

# def remove_run(run, p):
#     i = len(p.runs) - 1
#     while i >= 0:
#         if p.runs[i]._r == run._r:
#             print(run._r)
#             p._p.remove(p.runs[i]._r)
#             return run
#         i -= 1
#     return None

# def append_run(run, p):
#     r_copy = copy.deepcopy(run)._r
#     p._p.append(r_copy)

# def add_bullet_item(p, run):
#     p_new = utils.insert_paragraph_after(p, "", "List Bullet")
#     r_new = copy.deepcopy(run)._r
#     p_new._p.append(r_new)
#     return p_new

def duplicate(p):
    p_new = copy.deepcopy(p)
    p._p.addnext(p_new._p)
    return p_new

# def remove_runs_to_reverse(p, n):
#     i = len(p.runs) - 1
#     while i > n:
#         p._p.remove(p.runs[i]._r)
#         i -= 1

# def remove_runs_to(p, n):
#     i = 0
#     for i in range(n):
#         p._p.remove(p.runs[0]._r)